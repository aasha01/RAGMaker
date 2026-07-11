"""Manual parser — direct, library-per-format text extraction."""

from __future__ import annotations

import os

from .base import BaseParser, ParsedDocument


class ManualParser(BaseParser):
    """Extract text by calling a format-specific library directly, no framework.

    What it does (mechanically): dispatches on the file extension and reads the
    text with the most direct tool for that format — a plain file read for
    ``.txt``/``.md``, ``pypdf`` page-by-page for ``.pdf``, and ``python-docx``
    paragraph-by-paragraph for ``.docx``. It returns exactly what the library
    gives back, with no de-hyphenation, whitespace normalisation, or header/
    footer stripping. What you see is what the downstream stages get.

    Tradeoff vs. the alternatives: the LangChain and LlamaIndex loaders wrap
    these same libraries but add convenience (auto format detection, light
    cleanup, splitting into their own document objects). That convenience also
    hides decisions from you. This parser trades polish for a fully transparent,
    dependency-light path where every transformation is visible in the code.

    When a learner would prefer it: when you want to *see* the raw quality of an
    extractor — e.g. how a PDF's columns or tables come out as messy text —
    before deciding whether a fancier loader is worth it. It is also the
    zero-extra-framework default, so it is the right first thing to reach for.

    Per-format libraries are imported lazily inside ``parse`` so that, for
    example, parsing a ``.txt`` file needs no PDF or DOCX package installed.
    """

    name = "Manual (direct library)"
    description = (
        "Reads text with the most direct tool per format (plain read for txt, "
        "pypdf for PDF, python-docx for DOCX). No cleanup, no framework — the "
        "most transparent option. Best when you want to see exactly what an "
        "extractor produces before adding any convenience layer."
    )

    #: Extensions this parser knows how to handle, for a clear up-front error.
    SUPPORTED = (".txt", ".md", ".pdf", ".docx")

    def parse(self, file_path: str) -> ParsedDocument:
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"No such file to parse: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext in (".txt", ".md"):
            return self._parse_text(file_path, ext)
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        if ext == ".docx":
            return self._parse_docx(file_path)

        raise ValueError(
            f"ManualParser does not support '{ext}' files. "
            f"Supported: {', '.join(self.SUPPORTED)}. "
            f"(No silent fallback — pick a parser that handles this format.)"
        )

    def _parse_text(self, file_path: str, ext: str) -> ParsedDocument:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return ParsedDocument(
            text=text,
            source=os.path.basename(file_path),
            metadata={
                "format": ext.lstrip("."),
                "pages": 1,
                "page_texts": [text],
                "char_count": len(text),
            },
        )

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as e:  # pragma: no cover - environment dependent
            raise ImportError(
                "Parsing PDF with ManualParser needs the 'pypdf' package. "
                "Install it with: pip install pypdf"
            ) from e

        reader = PdfReader(file_path)
        page_texts = [(page.extract_text() or "") for page in reader.pages]
        # Join pages with a form-feed so the page boundaries stay visible/
        # recoverable rather than being silently smoothed away.
        text = "\f".join(page_texts)
        return ParsedDocument(
            text=text,
            source=os.path.basename(file_path),
            metadata={
                "format": "pdf",
                "pages": len(page_texts),
                "page_texts": page_texts,
                "char_count": len(text),
                "engine": "pypdf",
            },
        )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        try:
            import docx  # python-docx
        except ImportError as e:  # pragma: no cover - environment dependent
            raise ImportError(
                "Parsing DOCX with ManualParser needs the 'python-docx' package. "
                "Install it with: pip install python-docx"
            ) from e

        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs]
        text = "\n".join(paragraphs)
        return ParsedDocument(
            text=text,
            source=os.path.basename(file_path),
            metadata={
                "format": "docx",
                "paragraphs": len(paragraphs),
                "pages": 1,  # DOCX has no fixed page model without rendering
                "page_texts": [text],
                "char_count": len(text),
                "engine": "python-docx",
            },
        )
